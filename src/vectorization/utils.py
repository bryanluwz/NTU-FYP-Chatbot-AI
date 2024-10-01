import pandas as pd
from docx import Document


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


def extract_tables_from_docx(file_path):
    doc = Document(file_path)
    tables_data = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables_data.append(table_data)
    return tables_data


def extract_table_from_csv(file_path):
    df = pd.read_csv(file_path)
    return df.to_dict(orient='records')


def extract_table_from_excel(file_path):
    df = pd.read_excel(file_path)
    return df.to_dict(orient='records')
